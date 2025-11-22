from docx import Document
from docx2pdf import convert
from pathlib import Path
import sys
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# -------------------------------
# Helpers
# -------------------------------

def copy_run_attrs(src_run, dst_run):
    """Copia atributos de formatação de um run para outro."""
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb
    
    # Copiar estilo se houver
    dst_run.style = src_run.style

def set_cell_background(cell, color_hex):
    """Define a cor de fundo de uma célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def format_decimal(val):
    """Formata float para string com vírgula decimal (ex: 0.69 -> 0,69)."""
    if isinstance(val, (int, float)):
        return f"{val:.2f}".replace('.', ',')
    return str(val)

def format_currency(val):
    """Formata float para moeda BRL (ex: 199.0 -> R$ 199,00)."""
    if isinstance(val, (int, float)):
        # Formata com separador de milhar ponto e decimal vírgula
        s = f"{val:,.2f}"
        s = s.replace(',', 'X').replace('.', ',').replace('X', '.')
        return f"R$ {s}"
    return str(val)

def replace_placeholder_with_table(doc, placeholder, headers, row_label, data_values):
    """
    Substitui um parágrafo contendo o placeholder por uma tabela dinâmica.
    """
    target_paragraph = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_paragraph = p
            break
    
    if not target_paragraph:
        return

    # Criar tabela
    # Colunas = Label + Dados (len(data_values))
    cols = 1 + len(data_values)
    table = doc.add_table(rows=2, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Preencher Cabeçalho
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = headers[0]
    set_cell_background(hdr_cells[0], "D9D9D9") # Cinza claro
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, header_text in enumerate(headers[1:]):
        cell = hdr_cells[i+1]
        cell.text = str(header_text)
        set_cell_background(cell, "D9D9D9")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Preencher Dados
    row_cells = table.rows[1].cells
    row_cells[0].text = row_label
    # set_cell_background(row_cells[0], "FFFFFF") # Branco (padrão)
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, val in enumerate(data_values):
        cell = row_cells[i+1]
        cell.text = str(val)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Mover tabela para depois do parágrafo alvo
    target_paragraph._p.addnext(table._tbl)
    
    # Remover o parágrafo do placeholder
    p = target_paragraph._p
    p.getparent().remove(p)


# -------------------------------
# 1) Substituir placeholders
# -------------------------------

def substituir_placeholders_paragrafo(paragraph, mapa):
    # Loop para tratar placeholders um por um, evitando conflitos de índices
    while True:
        runs = paragraph.runs
        if not runs:
            break

        # 1. Tentar substituição direta em runs individuais (rápido e seguro)
        changed_any = False
        for run in runs:
            for chave, valor in mapa.items():
                if chave in run.text:
                    run.text = run.text.replace(chave, valor)
                    changed_any = True
        
        # Se mudou algo, recomeça para garantir consistência
        if changed_any:
            continue

        # 2. Verificar se ainda existem placeholders quebrados entre runs
        full_text = "".join(run.text for run in runs)
        found_key = None
        for k in mapa:
            if k in full_text:
                found_key = k
                break
        
        if not found_key:
            break # Nenhum placeholder restante

        # 3. Smart Merge: Identificar runs envolvidos APENAS nesse placeholder
        # Mapear cada caractere do full_text para o índice do run correspondente
        char_to_run_idx = []
        for i, run in enumerate(runs):
            char_to_run_idx.extend([i] * len(run.text))

        start_idx = full_text.find(found_key)
        end_idx = start_idx + len(found_key)
        
        # Identificar índices dos runs únicos envolvidos
        involved_run_indices = sorted(list(set(char_to_run_idx[start_idx:end_idx])))
        
        # Se por algum motivo não achou runs (ex: texto vazio?), aborta para evitar loop
        if not involved_run_indices:
            break

        # Merge: Juntar texto dos runs envolvidos no primeiro run
        first_run_idx = involved_run_indices[0]
        target_run = runs[first_run_idx]
        
        # Construir texto combinado desses runs
        merged_text = "".join(runs[i].text for i in involved_run_indices)
        
        # Substituir APENAS a primeira ocorrência da chave (a que encontramos)
        new_merged_text = merged_text.replace(found_key, mapa[found_key], 1)
        
        # Atualizar runs
        target_run.text = new_merged_text
        for i in involved_run_indices[1:]:
            runs[i].text = ""
            
        # Loop continua para processar próximos placeholders


def substituir_placeholders(documento, mapa):
    for paragraph in documento.paragraphs:
        substituir_placeholders_paragrafo(paragraph, mapa)

    for table in documento.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    substituir_placeholders_paragrafo(paragraph, mapa)


# -------------------------------
# 2) Remover negrito após rótulos
# -------------------------------

rotulos = [
    "Compradora:",
    "Ponto de Entrega:",
    "Período de fornecimento:", 
    "Período de fornecimento",
    "Fonte de Fornecimento:",
    "Obs.:",
    "flexibilidade mensal:",
    "sazonalidade:",
    "Modulação:",
    "Faturamento e pagamento:",
    "Garantia de pagamento:",
    "Reajuste:"
]


def remover_negrito_apos_rotulos(documento):

    def process_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        
        # Verifica se o parágrafo contém algum dos rótulos
        matched_label = None
        match_index = -1
        
        for label in rotulos:
            if label in full_text:
                idx = full_text.find(label)
                if idx != -1:
                    matched_label = label
                    match_index = idx
                    break
        
        if not matched_label:
            return

        split_point = match_index + len(matched_label)
        
        if split_point >= len(full_text):
            return

        # Coletar dados dos runs para reconstrução preservando formatação
        runs_data = []
        current_pos = 0
        
        for run in paragraph.runs:
            run_len = len(run.text)
            run_end = current_pos + run_len
            
            # Caso 1: Run totalmente antes do split point (Manter como está)
            if run_end <= split_point:
                runs_data.append({
                    "text": run.text,
                    "src_run": run,
                    "force_unbold": False
                })
            
            # Caso 2: Run totalmente depois do split point (Forçar sem negrito)
            elif current_pos >= split_point:
                runs_data.append({
                    "text": run.text,
                    "src_run": run,
                    "force_unbold": True
                })
                
            # Caso 3: Run atravessa o split point (Dividir)
            else:
                split_in_run = split_point - current_pos
                text_before = run.text[:split_in_run]
                text_after = run.text[split_in_run:]
                
                if text_before:
                    runs_data.append({
                        "text": text_before,
                        "src_run": run,
                        "force_unbold": False
                    })
                if text_after:
                    runs_data.append({
                        "text": text_after,
                        "src_run": run,
                        "force_unbold": True
                    })
            
            current_pos += run_len

        # Limpar runs existentes
        p = paragraph._p
        for run in paragraph.runs:
            p.remove(run._r)
            
        # Reconstruir runs
        for data in runs_data:
            new_run = paragraph.add_run(data["text"])
            copy_run_attrs(data["src_run"], new_run)
            
            if data["force_unbold"]:
                new_run.bold = False

    # Texto normal
    for paragraph in documento.paragraphs:
        process_paragraph(paragraph)

    # Também processa texto interno de tabelas
    for table in documento.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)


# ==========================================================
# 3) GERAÇÃO DO DOCUMENTO FINAL
# ==========================================================

def generate_proposal(
    data_hoje,
    razao_social,
    cnpj,
    submercado,
    inicio,
    fim,
    curva_vol,
    curva_precos,
    anos,
    tipo_energia,
    flex,
    sazo,
    modulacao,
    pagamento,
    qty_meses,
    tipo_proposta
):
    # Mapa de valores simples
    valores = {
        "{{DATA_HOJE}}": data_hoje,
        "{{RAZAO_SOCIAL}}": razao_social,
        "{{CNPJ}}": cnpj,
        "{{SUBMERCADO}}": submercado,
        "{{INICIO}}": inicio,
        "{{FIM}}": fim,
        # "{{CURVA_VOL}}": Removido, tratado via tabela
        # "{{CURVA_PRECOS}}": Removido, tratado via tabela
        "{{TIPO_ENERGIA}}": tipo_energia,
        "{{FLEX}}": flex,
        "{{SAZO}}": sazo,
        "{{MODULACAO}}": modulacao,
        "{{PAGAMENTO}}": pagamento,
        "{{QTY_MESES}}": qty_meses,
        "{{TIPO_PROPOSTA}}": tipo_proposta,
    }

    BASE_DIR = Path(__file__).resolve().parent.parent
    arquivo_origem = BASE_DIR / "assets" / "documents" / "standard_proposal.docx"
    arquivo_final_docx = BASE_DIR / "assets" / "documents" / "standard_proposal_preenchida.docx"
    arquivo_final_pdf = BASE_DIR / "assets" / "documents" / "standard_proposal_preenchida.pdf"

    print(f"Iniciando geração de proposta...")

    try:
        doc = Document(arquivo_origem)

        # 1. Criar tabelas dinâmicas (Volume e Preço)
        # Preparar dados formatados
        headers = ["Ano"] + [str(a) for a in anos]
        
        # Tabela Volume
        vol_values = [format_decimal(v) for v in curva_vol]
        replace_placeholder_with_table(doc, "{{CURVA_VOL}}", headers, "Vol (MWm)", vol_values)
        
        # Tabela Preço
        price_values = [format_currency(p) for p in curva_precos]
        replace_placeholder_with_table(doc, "{{CURVA_PRECOS}}", headers, "Preço", price_values)

        # 2. Substituir placeholders restantes
        substituir_placeholders(doc, valores)
        
        # 3. Remover negrito após rótulos
        remover_negrito_apos_rotulos(doc)
        
        # Salvar e Converter
        doc.save(arquivo_final_docx)
        print("Documento DOCX salvo.")

        convert(str(arquivo_final_docx), str(arquivo_final_pdf))
        print("PDF gerado com sucesso.")

        return str(arquivo_final_pdf)

    except PermissionError:
        print("ERRO: O arquivo está aberto em outro programa. Por favor, feche o Word e tente novamente.")
        raise
    except Exception as e:
        print(f"ERRO INESPERADO: {e}")
        raise


if __name__ == "__main__":
    # Teste manual
    generate_proposal(
        data_hoje="21/11/2025",
        razao_social="flex IMPORT COMERCIO E INDUSTRIA LTDA",
        cnpj="08.297.453/0001-33",
        submercado="Nordeste",
        inicio="01/01/2026",
        fim="31/12/2028",
        curva_vol=[0.69, 0.69, 0.69, 0.5, 0.4],
        curva_precos=[199, 199, 197, 200, 180],
        anos=[2026, 2027, 2028, 2029, 2030],
        tipo_energia="Energia Incentivada 50%",
        flex="30",
        sazo="100",
        modulacao="Flat",
        pagamento="6",
        qty_meses="2",
        tipo_proposta="Oficial"
    )
